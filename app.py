"""
Project: AI Chatbot
Date: 03/25
Author: James W.
Desc: CodeLlama-34b-Instruct-hf knowledge cutoff is December 31, 2022. Maximum context window (e.g., 4096 tokens)
"""

import gradio as gr
import logging
import os
from dotenv import load_dotenv
from huggingface_hub import InferenceClient
from docx import Document

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

load_dotenv()
api_key = os.getenv("HUGGINGFACE_API_KEY") #get your huggingface api key and put it in .env

client = InferenceClient(
    #base_url="https://router.huggingface.co/hf-inference/models/codellama/CodeLlama-34b-Instruct-hf/v1",
    provider="hf-inference",#default
    api_key=api_key
)

"""history_msg:
[
    {"role": "user", "content": "What is the capital of France?"},
    {"role": "assistant", "content": "Paris"}
]
if file is not None:
    if file.endswith(".txt"):
        with open(file, "r", encoding="utf-8") as f:
            file_content = f.read()
    elif file.endswith(".docx"):  # Note: .doc might need a different library
        doc = Document(file)
        file_content = "\n".join([para.text for para in doc.paragraphs])
    system_msg += f" Additional context from uploaded file: {file_content}"
"""

def respond(current_msg,
            history_msg,
            max_tokens,
            temperature,
            file):
    system_msg = "You are a friendly assistant chatbot. Answer directly and concisely."
    if file is not None:
        if file.endswith(".txt"):
            with open(file, "r", encoding="utf-8") as f:
                file_content=f.read()
        elif file.endswith(".docx"):
            doc=Document(file)
            #file_content="\n".join([para.text for para in doc.paragraphs])
            file_content=""
            for paragraph in doc.paragraphs:
                file_content+=paragraph.text + "\n"            
        system_msg+= f" Additional context from uploaded file:{file_content}"
    messages = [{"role": "system", "content": system_msg}]
    for msg in history_msg:
        messages.append(
            {"role": msg.get("role"), "content": msg.get("content")})
    messages.append({"role": "user", "content": current_msg})
    response = ""
    try:
        chat_completion_output = client.chat.completions.create(
            messages=messages,
            model="codellama/CodeLlama-34b-Instruct-hf",
            max_tokens=max_tokens,
            temperature=temperature,
            stream=False
        )
        response = chat_completion_output.choices[0].message.content
    except ConnectionError as e:
        logger.error(f"Network error: {e}")
        return None
    except ValueError as e:
        logger.error(f"Invalid input: {e}")
        return None
    except Exception as e:
        logger.error(f"Unexpected error: {e}")
        return None
    else:
        return response
    finally:
        logger.info("Execution completed.")


chatbot = gr.ChatInterface(fn=respond,
                        type="messages",
                        additional_inputs=[
                            gr.Slider(minimum=1, maximum=2048, value=256,
                                      step=1, label="Max output tokens"),
                            gr.Slider(minimum=0.1, maximum=1.0, value=0.2, step=0.1, label="Creativeness"),
                            gr.File(label="Upload a text file", file_types=[".txt",".docx"])
                            ])

#if __name__ == "__main__":
chatbot.launch(share=True)


# import gradio as gr
# from huggingface_hub import InferenceClient

# """
# For more information on `huggingface_hub` Inference API support, please check the docs: https://huggingface.co/docs/huggingface_hub/v0.22.2/en/guides/inference
# """
# client = InferenceClient(
#     #"deepseek-ai/DeepSeek-R1-Distill-Qwen-1.5B",
#     api_key=api_key,
#     model= "codellama/CodeLlama-34b-Instruct-hf"
#     #provider="nebius"
#     #provider="fireworks-ai"
#     )
# def respond(
#     message,
#     history,
#     system_message,
#     max_tokens,
#     temperature,
#     top_p,
# ):
#     # Clean history to keep only "role" and "content" his error occurs when using the Hugging Face Inference API (via InferenceClient). The messages list in your request includes extra fields (metadata and options) that the API doesn’t allow. These fields are likely coming from Gradio’s ChatInterface with type="messages", which adds them automatically
#     history_cleaned = [{"role": msg["role"], "content": msg["content"]} for msg in history]
    
#     # Build the messages list
#     messages = [{"role": "system", "content": system_message}] + history_cleaned
#     messages.append({"role": "user", "content": message})

#     try:
#         completion = client.chat_completion(
#             messages,
#             #model="perplexity-ai/r1-1776",
#             max_tokens=max_tokens,
#             stream=False,
#             temperature=temperature,
#             top_p=top_p,
#         )
#         response = completion.choices[0].message.content
#         print(response)
#         return response
#     except Exception as e:
#         print(f"My Error: {e}")
#         return f"My Error: {e}"
    
# """
# For information on how to customize the ChatInterface, peruse the gradio docs: https://www.gradio.app/docs/chatinterface
# """
# demo = gr.ChatInterface(
#     respond, #it's not respond(), not call the funcion?
#     type="messages",
#     additional_inputs=[
#         gr.Textbox(value="You are a friendly chatbot. Answer directly and concisely, no reasoning or explanations.", label="System message"),
#         gr.Slider(minimum=1, maximum=2048, value=64, step=1, label="Max new tokens"),
#         gr.Slider(minimum=0.1, maximum=4.0, value=0.7, step=0.1, label="Temperature"),
#         gr.Slider(
#             minimum=0.1,
#             maximum=1.0,
#             value=0.95,
#             step=0.05,
#             label="Top-p (nucleus sampling)",
#         ),
#     ],
# )


# if __name__ == "__main__": 
#     demo.launch()
#如果脚本直接运行（而非被导入），启动 Gradio 界面。if __name__ == "__main__": 确保某些代码（这里是 demo.launch()）只在脚本被直接运行时执行，而不是在被导入时意外触发。

# try:
#     for message in client.chat_completion(...):
#         token = message.choices[0].delta.content or ""
#         response += token
#         yield response
# except Exception as e:
#     yield f"错误: {e}"


# from huggingface_hub import InferenceClient

# # 使用你的 Hugging Face token
# client = InferenceClient(api_key="")  # 替换为真实 token

# # 发送请求
# messages = [{"role": "user", "content": "法国的首都是哪里？"}]
# try:
#     completion = client.chat.completions.create(
#         model="google/gemma-2-2b-it",  # 确认此模型支持 API
#         messages=messages,
#         max_tokens=500
#     )
#     print(completion.choices[0].message.content)
# except Exception as e:
#     print(f"错误: {e}")

# import gradio as gr
# from huggingface_hub import InferenceClient

# client = InferenceClient(api_key=api_key)

# def chat_function(user_input, chat_history):
#     messages = [{"role": "user", "content": user_input}]
#     response = client.chat.completions.create(
#         model="google/gemma-2-2b-it",
#         messages=messages,
#         max_tokens=500
#     )
#     chat_history.append({"role": "user", "content": user_input})
#     chat_history.append({"role": "assistant", "content": response.choices[0].message.content})
#     return chat_history

# with gr.Blocks() as demo:
#     gr.Chatbot(label="HF Chat", type="messages")
#     gr.Textbox(placeholder="输入消息...", label="输入").submit(chat_function, [gr.Textbox(), gr.Chatbot()], gr.Chatbot())

# demo.launch(server_name="0.0.0.0", server_port=7860)